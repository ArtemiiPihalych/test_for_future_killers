from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


SOURCE = Path("gosy_test_2.docx")
OUT = Path("data/questions.private.json")
JS_OUTS = [Path("questions.js"), Path("public/questions.js")]


def clean(text: str) -> str:
    text = " ".join(text.replace("\u00a0", " ").split())
    text = re.sub(r"\s*[-–]?\s*(?:это ключи|так в ключах|это оно)\s*", "", text, flags=re.IGNORECASE)
    text = text.replace("\uf00c", "")
    text = re.sub(
        r"\s*\([^)]*(?:гпт|нейронк|не ебу|хз|карпова|ключ)[^)]*\)",
        "",
        text,
        flags=re.IGNORECASE,
    )
    if any(marker in text for marker in ("Р", "С", "вЂ", "О±", "ОІ", "Оі")):
        try:
            repaired = text.encode("cp1251").decode("utf-8")
            if sum("а" <= ch.lower() <= "я" for ch in repaired) > sum(
                "а" <= ch.lower() <= "я" for ch in text
            ):
                return repaired
        except UnicodeError:
            pass
    return text


def bold_key(paragraph) -> str:
    text = "".join(run.text for run in paragraph.runs if run.bold)
    return "".join(ch for ch in text if ch.isalnum())


def line_from_paragraph(paragraph) -> dict | None:
    text = clean(paragraph.text)
    if not text:
        return None
    return {"text": text, "bold": bold_key(paragraph)}


def line_from_cell(cell) -> dict | None:
    text = clean(cell.text.replace("\n", " "))
    if not text:
        return None
    bold = "".join(
        run.text
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.bold
    )
    return {"text": text, "bold": "".join(ch for ch in bold if ch.isalnum())}


def has_meaningful_bold(line: dict) -> bool:
    key = line["bold"]
    text = line["text"]
    if not key:
        return False
    option_prefix = re.match(r"^\s*(?:[0-9]+|[a-zA-Zа-яА-Я])[\).]?\s*", text)
    if option_prefix:
        prefix_key = "".join(ch for ch in option_prefix.group(0) if ch.isalnum())
        key = key[len(prefix_key) :] if key.startswith(prefix_key) else key
    return len(key) >= 1 and any(ch.isalnum() for ch in key)


def is_title(line: dict) -> bool:
    return "Тестовые задания по дисциплине" in line["text"]


def is_carry_prompt(line: dict) -> bool:
    text = line["text"]
    return text.endswith(":") and not has_meaningful_bold(line)


def make_question(lines: list[dict], carry: list[str] | None = None, option_count: int = 4) -> dict | None:
    if len(lines) < option_count + 1:
        return None
    prompt_lines = lines[:-option_count]
    option_lines = lines[-option_count:]
    prompt = " ".join([*(carry or []), *(line["text"] for line in prompt_lines)])
    options = [line["text"] for line in option_lines]
    correct = [idx for idx, line in enumerate(option_lines) if has_meaningful_bold(line)]
    if not prompt or len(options) != option_count or len(correct) != 1:
        return None
    return {"question": prompt, "options": options, "answer": correct[0]}


def make_manual_question(prompt: str, options: list[str], answer: int) -> dict:
    return {"question": clean(prompt), "options": [clean(option) for option in options], "answer": answer}


def manual_block(index: int, block: list[dict]) -> list[dict] | None:
    simple_answers = {
        70: 0,
        78: 1,
        79: 3,
        80: 0,
        87: 3,
        104: 3,
        117: 3,
        294: 1,
        295: 1,
        298: 2,
        299: 1,
        300: 0,
        301: 2,
        304: 0,
        305: 0,
        307: 0,
        309: 2,
        310: 1,
        311: 3,
        312: 0,
        313: 2,
        314: 0,
        315: 3,
        316: 3,
        317: 3,
        318: 2,
        319: 0,
        320: 1,
        321: 2,
        322: 0,
        323: 0,
        324: 3,
        325: 1,
        326: 0,
        327: 1,
        328: 0,
        329: 0,
        396: 2,
        475: 1,
        552: 0,
        553: 0,
        622: 1,
        1017: 0,
        1037: 0,
        1039: 3,
        1048: 1,
        1065: 0,
        1069: 0,
        1073: 0,
        1134: 2,
        1358: 2,
    }
    lines = [line["text"] for line in block]
    if index in simple_answers:
        answer = simple_answers[index]
        if index == 1065:
            return [make_manual_question(lines[0], [lines[1], lines[2], f"{lines[3]} {lines[4]}", lines[5]], answer)]
        if index == 1069:
            return [make_manual_question(lines[0], [f"{lines[1]} {lines[2]}", lines[3], lines[4], f"{lines[5]} {lines[6]}"], answer)]
        if index == 1073:
            return [make_manual_question(lines[0], [f"{lines[1]} {lines[2]}", f"{lines[3]} {lines[4]}", lines[5], lines[6]], answer)]
        if index == 1358:
            return [make_manual_question(lines[0], [lines[2], lines[3], f"{lines[4]} {lines[5]}", lines[6]], answer)]
        return [make_manual_question(lines[0], lines[1:], answer)]

    if index == 219:
        return [
            make_manual_question(
                "1. ЛОГИЧЕСКИЙ ПРИЕМ, МЕТОД ИССЛЕДОВАНИЯ, ОЗНАЧАЮЩИЙ МЫСЛЕННОЕ РАЗЛОЖЕНИЕ ОБЪЕКТА НА СОСТАВНЫЕ ЭЛЕМЕНТЫ НАЗЫВАЕТСЯ:",
                ["1. экстраполяцией", "2. синтезом", "3. аналогией", "4. анализом"],
                3,
            )
        ]
    if index == 220:
        return [make_manual_question(lines[0], ["1. эмпиризм", "2. материализм", "3. идеализм", "4. гедонизм"], 0)]
    if index == 308:
        return [make_manual_question(lines[0], lines[1:5], 1), make_manual_question(lines[5], lines[6:10], 2)]
    if index == 340:
        return [
            make_manual_question(
                "10.Клон лимфоцитов – это:",
                [
                    "1.Потомство одной клетки, отличающееся по специфичности рецепторов",
                    "2.Группа всех лимфоцитов",
                    "3.Потомство разных клеток",
                    "4.Группа лейкоцитов",
                ],
                0,
            )
        ]
    if index == 1092:
        return [make_manual_question(lines[0], lines[1:5], 0), make_manual_question(lines[5], lines[6:10], 0)]
    return None


def split_block(block: list[dict], carry: list[str] | None) -> tuple[list[dict], list[str] | None]:
    questions: list[dict] = []
    local_carry = carry

    if len(block) == 1 and is_title(block[0]):
        return questions, None

    if len(block) <= 2 and all(is_carry_prompt(line) or not has_meaningful_bold(line) for line in block):
        return questions, [*(local_carry or []), *(line["text"] for line in block)]

    if len(block) == 4 and local_carry:
        question = make_question([{"text": " ".join(local_carry), "bold": ""}, *block])
        return ([question] if question else []), None

    if len(block) == 4:
        question = make_question(block, local_carry, option_count=3)
        return ([question] if question else []), None

    if len(block) == 6 and is_carry_prompt(block[-1]):
        question = make_question(block[:5], local_carry)
        if question:
            questions.append(question)
        return questions, [block[-1]["text"]]

    if len(block) == 6:
        question = make_question(block, local_carry)
        return ([question] if question else []), None

    if len(block) >= 10 and len(block) % 5 == 0:
        for idx in range(0, len(block), 5):
            question = make_question(block[idx : idx + 5], local_carry)
            if question:
                questions.append(question)
        return questions, None

    if len(block) >= 11 and len(block) % 5 == 1 and is_carry_prompt(block[0]):
        header = [*(local_carry or []), block[0]["text"]]
        for idx in range(1, len(block), 5):
            question = make_question(block[idx : idx + 5], header)
            if question:
                questions.append(question)
        return questions, None

    question = make_question(block, local_carry)
    return ([question] if question else []), None


def main() -> None:
    document = Document(SOURCE)
    blocks: list[list[dict]] = []
    current: list[dict] = []

    for paragraph in document.paragraphs:
        line = line_from_paragraph(paragraph)
        if line is None:
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line)

    if current:
        blocks.append(current)

    for table in document.tables:
        current = []
        for row in table.rows:
            line = line_from_cell(row.cells[0])
            if line is None:
                if current:
                    blocks.append(current)
                    current = []
                continue
            current.append(line)
        if current:
            blocks.append(current)

    questions: list[dict] = []
    unresolved: list[dict] = []
    carry: list[str] | None = None

    for index, block in enumerate(blocks):
        manual = manual_block(index, block)
        if manual is not None:
            questions.extend(manual)
            carry = None
            continue

        parsed, carry = split_block(block, carry)
        if parsed:
            questions.extend(parsed)
        elif not carry and not (len(block) == 1 and is_title(block[0])):
            unresolved.append({"block": index, "lines": [line["text"] for line in block]})

    deduped: list[dict] = []
    seen: set[tuple[str, tuple[str, ...]]] = set()
    for idx, question in enumerate(questions, start=1):
        key = (question["question"], tuple(question["options"]))
        if key in seen:
            continue
        seen.add(key)
        question["id"] = idx
        deduped.append(question)

    OUT.parent.mkdir(exist_ok=True)
    json_payload = json.dumps(deduped, ensure_ascii=False, indent=2)
    OUT.write_text(json_payload, encoding="utf-8")
    js_payload = f"window.QUESTION_BANK = {json_payload};\n"
    for path in JS_OUTS:
        path.parent.mkdir(exist_ok=True)
        path.write_text(js_payload, encoding="utf-8")
    Path("data/extraction-report.json").write_text(
        json.dumps(
            {
                "source": str(SOURCE),
                "questions": len(deduped),
                "unresolved_blocks": unresolved,
                "unresolved_count": len(unresolved),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(f"questions={len(deduped)} unresolved={len(unresolved)}")


if __name__ == "__main__":
    main()
